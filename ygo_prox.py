import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO

HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}

def get_card_image(card_name: str) -> str | None:
    """
    Search DeviantArt for 'rush duel <card_name>',
    follow the first result, and return the full-size image URL.
    """
    query = f"rush duel {card_name}".replace(" ", "+")
    search_url = f"https://www.deviantart.com/search?q={query}"

    try:
        resp = requests.get(search_url, headers=HEADERS, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")

        # First artwork link in results
        first_link = soup.select_one("a[href][aria-label][aria-label*='by']")
        if not first_link:
            return None

        art_url = first_link["href"]

        # Open artwork page and grab full-size image
        art_resp = requests.get(art_url, headers=HEADERS, timeout=10)
        art_resp.raise_for_status()
        art_soup = BeautifulSoup(art_resp.text, "html.parser")

        img_tag = art_soup.select_one("div[typeof=ImageObject] img")
        if img_tag and img_tag.get("src"):
            return img_tag["src"]

    except Exception as e:
        print(f"Error fetching {card_name}: {e}")

    return None


def decklist_to_docx(
    deck_file: str,
    output_docx: str,
    card_width_in=2.31,
    card_height_in=3.37,
    per_row=3
):
    """
    Reads a decklist, fetches card images from DeviantArt, and builds a Word document.
    Ensures paragraph (row) spacing is zero so cards sit flush vertically.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Parse decklist
    skip_headers = {
        "monster", "spell", "trap", "extra", "side",
        "몬스터", "마법", "함정", "엑스트라", "사이드"
    }
    entries: list[tuple[int, str]] = []

    with open(deck_file, "r", encoding="utf-8") as f:
        for raw in f:
            line = raw.strip()
            if not line:
                continue
            if line.lower() in skip_headers:
                continue

            parts = line.split(" ", 1)
            if len(parts) != 2:
                continue
            cnt, name = parts
            if not cnt.isdigit():
                continue
            entries.append((int(cnt), name))

    # Build document (rows of images)
    current_row_runs = 0
    paragraph = None

    for count, card_name in entries:
        print(f"Searching: {card_name}")
        img_url = get_card_image(card_name)
        if not img_url:
            print(f"⚠️ No image found for {card_name}")
            continue

        # Download image once and reuse
        try:
            r = requests.get(img_url, headers=HEADERS, timeout=15)
            r.raise_for_status()
            img_blob = BytesIO(r.content)
        except Exception as e:
            print(f"Error downloading {card_name}: {e}")
            continue

        # Insert as many copies as needed
        for _ in range(count):
            # start a new row if needed
            if paragraph is None or current_row_runs >= per_row:
                paragraph = doc.add_paragraph()

                # 🔻 Set tight spacing for this row
                pf = paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0

                current_row_runs = 0

            run = paragraph.add_run()

            # Rewind buffer for each insertion
            img_blob.seek(0)
            run.add_picture(img_blob, width=Inches(card_width_in), height=Inches(card_height_in))
            current_row_runs += 1

    doc.save(output_docx)
    print(f"Done! Saved: {output_docx}")


if __name__ == "__main__":
    decklist_to_docx("deck.txt", "deck_proxies.docx", per_row=3)
